"""Generate SAB India Tracker flow chart Word document.

Produces docs/SAB_India_Tracker_Flowcharts.docx with eight flow diagrams rendered
as PNG via matplotlib and embedded into a structured Word document.
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = Path(__file__).parent
IMG_DIR = DOCS_DIR / "flowchart_images"
IMG_DIR.mkdir(exist_ok=True)

# ── Colour palette (paper / signal-orange industrial) ─────────────────────────
C_INK = "#2b2624"
C_INK3 = "#8a7f78"
C_PAPER = "#f9f6f1"
C_ACCENT = "#d4712b"
C_ACCENT_WASH = "#fbeadb"
C_BLUE = "#3c7ab8"
C_BLUE_WASH = "#e2ecf5"
C_POSITIVE = "#4a8c5e"
C_POSITIVE_WASH = "#dcefe2"
C_AMBER = "#c98a27"
C_AMBER_WASH = "#f7ecd2"
C_ALERT = "#c14330"
C_ALERT_WASH = "#f5d9d3"
C_RULE = "#cbc4bd"


def _box(ax, xy, w, h, text, fill=C_PAPER, stroke=C_RULE, text_color=C_INK,
         bold=False, fontsize=9):
    x, y = xy
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.12",
                         linewidth=1.1, edgecolor=stroke, facecolor=fill)
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fontsize, color=text_color,
            fontweight="bold" if bold else "normal", wrap=True)


def _diamond(ax, xy, w, h, text, fill=C_AMBER_WASH, stroke=C_AMBER,
             text_color=C_INK, fontsize=9):
    cx, cy = xy[0] + w / 2, xy[1] + h / 2
    pts = [(cx, xy[1] + h), (xy[0] + w, cy), (cx, xy[1]), (xy[0], cy)]
    d = patches.Polygon(pts, closed=True, linewidth=1.1,
                        edgecolor=stroke, facecolor=fill)
    ax.add_patch(d)
    ax.text(cx, cy, text, ha="center", va="center",
            fontsize=fontsize, color=text_color, fontweight="bold")


def _arrow(ax, start, end, label=None, color=C_INK, lw=1.2):
    a = FancyArrowPatch(start, end, arrowstyle="-|>",
                        mutation_scale=14, linewidth=lw, color=color,
                        connectionstyle="arc3,rad=0")
    ax.add_patch(a)
    if label:
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        ax.text(mx, my + 0.12, label, ha="center", va="bottom",
                fontsize=8, color=C_INK3, fontstyle="italic")


def _frame(title: str, w=12, h=7, figsize=None):
    if figsize is None:
        figsize = (w, h)
    fig, ax = plt.subplots(figsize=figsize, dpi=150)
    ax.set_xlim(0, w)
    ax.set_ylim(0, h)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("white")
    ax.text(0.1, h - 0.25, title, fontsize=13, fontweight="bold", color=C_INK)
    ax.plot([0.1, w - 0.1], [h - 0.45, h - 0.45], color=C_RULE, lw=0.6)
    return fig, ax


def _save(fig, name: str) -> Path:
    path = IMG_DIR / f"{name}.png"
    fig.savefig(path, bbox_inches="tight", facecolor="white", dpi=150)
    plt.close(fig)
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 1. System navigation map
# ──────────────────────────────────────────────────────────────────────────────
def fig_nav_map():
    W, H = 12, 8
    fig, ax = _frame("1. System navigation — where everything lives", w=W, h=H)

    # Top row: login → desktop
    _box(ax, (0.5, 6.6), 2.4, 0.7, "Login\n(work email + SSO)",
         fill=C_ACCENT_WASH, stroke=C_ACCENT, bold=True)
    _box(ax, (3.5, 6.6), 2.6, 0.7, "Desktop App\nOperations Home",
         fill=C_INK, stroke=C_INK, text_color="white", bold=True)
    _arrow(ax, (2.9, 6.95), (3.5, 6.95))

    # Side-channels (mobile + public)
    _box(ax, (6.7, 6.6), 2.4, 0.7,
         "Mobile punch app\n(site crew)", fill=C_POSITIVE_WASH,
         stroke=C_POSITIVE, fontsize=9)
    _box(ax, (9.3, 6.6), 2.4, 0.7,
         "Public invoice share\n(client view)", fill=C_BLUE_WASH,
         stroke=C_BLUE, fontsize=9)

    # Section headers
    section_y = 5.4
    sections = [
        ("Sales", C_BLUE, C_BLUE_WASH, 0.5,
         ["Home", "Projects", "Clients", "Quotes", "Tax invoices"]),
        ("Operations", C_POSITIVE, C_POSITIVE_WASH, 3.4,
         ["Timesheets", "Inventory", "Overhead", "Reports"]),
        ("Procurement", C_ACCENT, C_ACCENT_WASH, 6.3,
         ["Vendors", "Purchase orders", "Goods receipts", "Vendor bills"]),
        ("Admin", C_AMBER, C_AMBER_WASH, 9.2,
         ["Users & roles", "Wage rates"]),
    ]
    for name, stroke, fill, x, items in sections:
        _box(ax, (x, section_y), 2.4, 0.55, name, fill=stroke, stroke=stroke,
             text_color="white", bold=True, fontsize=10)
        for i, item in enumerate(items):
            _box(ax, (x, section_y - 0.7 - i * 0.55), 2.4, 0.45, item,
                 fill=fill, stroke=stroke, fontsize=9)
        _arrow(ax, (4.8, 6.6), (x + 1.2, section_y + 0.55),
               color=C_INK3, lw=0.9)

    # Footer note pinned to bottom
    _box(ax, (0.5, 0.3), 11.2, 0.55,
         "Sign-in lands on Operations Home; the home action list pulls outstanding work from every other module.",
         fill=C_PAPER, stroke=C_RULE, fontsize=9)

    return _save(fig, "01_nav_map")


# ──────────────────────────────────────────────────────────────────────────────
# 2. Procurement — Vendor → PO → GRN → Bill → Payment
# ──────────────────────────────────────────────────────────────────────────────
def fig_procurement_flow():
    W, H = 14, 9
    fig, ax = _frame("2. Procurement — vendor to payment (3-way match)", w=W, h=H)

    # Row 1: Vendor → Draft PO → Decision diamond
    _box(ax, (0.3, 7.5), 2.6, 0.8,
         "Vendor master\nGSTIN · terms · MSME", fill=C_ACCENT_WASH,
         stroke=C_ACCENT, bold=True)
    _box(ax, (3.4, 7.5), 2.6, 0.8,
         "Draft PO\nvendor · project · lines",
         fill=C_PAPER, stroke=C_RULE)
    _arrow(ax, (2.9, 7.9), (3.4, 7.9))

    _diamond(ax, (6.5, 7.0), 3.0, 1.6,
             "Approval tier?\nbased on total",
             fill=C_AMBER_WASH, stroke=C_AMBER, fontsize=9)
    _arrow(ax, (6.0, 7.9), (6.5, 7.85))

    _box(ax, (10.2, 8.3), 3.4, 0.55, "≤ ₹1L → auto-approved",
         fill=C_POSITIVE_WASH, stroke=C_POSITIVE, bold=True, fontsize=9)
    _box(ax, (10.2, 7.65), 3.4, 0.55, "₹1–10L → PM (Anita Rao)",
         fill=C_BLUE_WASH, stroke=C_BLUE, bold=True, fontsize=9)
    _box(ax, (10.2, 7.0), 3.4, 0.55, "> ₹10L → Director",
         fill=C_ALERT_WASH, stroke=C_ALERT, bold=True, fontsize=9)
    _arrow(ax, (9.5, 8.0), (10.2, 8.55))
    _arrow(ax, (9.5, 7.85), (10.2, 7.9))
    _arrow(ax, (9.5, 7.7), (10.2, 7.25))

    # Row 2: PO sent → goods arrive → GRN
    _box(ax, (3.4, 5.6), 2.6, 0.7,
         "PO approved & sent",
         fill=C_BLUE_WASH, stroke=C_BLUE, bold=True)
    _arrow(ax, (11.9, 7.0), (4.7, 6.3))

    _box(ax, (0.3, 4.3), 2.6, 0.7,
         "Goods arrive at site",
         fill=C_PAPER, stroke=C_RULE, fontsize=9)
    _box(ax, (3.4, 4.3), 2.6, 0.7,
         "GRN recorded\nrecv · accepted · rejected",
         fill=C_ACCENT_WASH, stroke=C_ACCENT, bold=True, fontsize=9)
    _arrow(ax, (4.7, 5.6), (4.7, 5.0))
    _arrow(ax, (2.9, 4.65), (3.4, 4.65))

    _diamond(ax, (6.5, 3.8), 3.0, 1.6,
             "GRN status?",
             fill=C_AMBER_WASH, stroke=C_AMBER, fontsize=9)
    _arrow(ax, (6.0, 4.65), (6.5, 4.6))

    _box(ax, (10.2, 5.1), 3.4, 0.55, "all accepted → received",
         fill=C_POSITIVE_WASH, stroke=C_POSITIVE, fontsize=9)
    _box(ax, (10.2, 4.45), 3.4, 0.55, "partial → part-received",
         fill=C_AMBER_WASH, stroke=C_AMBER, fontsize=9)
    _box(ax, (10.2, 3.8), 3.4, 0.55, "rejected → credit note",
         fill=C_ALERT_WASH, stroke=C_ALERT, fontsize=9)
    _arrow(ax, (9.5, 4.7), (10.2, 5.37))
    _arrow(ax, (9.5, 4.6), (10.2, 4.72))
    _arrow(ax, (9.5, 4.5), (10.2, 4.07))

    # Row 3: Vendor bill → 3-way match
    _box(ax, (0.3, 2.5), 2.6, 0.7,
         "Vendor invoice in",
         fill=C_PAPER, stroke=C_RULE, fontsize=9)
    _box(ax, (3.4, 2.5), 2.6, 0.7,
         "Bill recorded\nsubtotal + GST split",
         fill=C_ACCENT_WASH, stroke=C_ACCENT, bold=True, fontsize=9)
    _arrow(ax, (2.9, 2.85), (3.4, 2.85))
    _arrow(ax, (4.7, 4.3), (4.7, 3.2))

    _diamond(ax, (6.5, 2.0), 3.0, 1.6,
             "3-way match\nPO ↔ GRN ↔ Bill",
             fill=C_BLUE_WASH, stroke=C_BLUE, fontsize=9)
    _arrow(ax, (6.0, 2.85), (6.5, 2.8))

    _box(ax, (10.2, 3.3), 3.4, 0.55, "matched → ready to pay",
         fill=C_POSITIVE_WASH, stroke=C_POSITIVE, fontsize=9)
    _box(ax, (10.2, 2.65), 3.4, 0.55, "discrepancy → buyer fix",
         fill=C_ALERT_WASH, stroke=C_ALERT, fontsize=9)
    _box(ax, (10.2, 2.0), 3.4, 0.55, "no PO → manual approve",
         fill=C_AMBER_WASH, stroke=C_AMBER, fontsize=9)
    _arrow(ax, (9.5, 2.9), (10.2, 3.57))
    _arrow(ax, (9.5, 2.8), (10.2, 2.92))
    _arrow(ax, (9.5, 2.7), (10.2, 2.27))

    # Bottom row: payment
    _box(ax, (3.4, 0.7), 3.0, 0.7,
         "Schedule payment\npick date · status → approved",
         fill=C_POSITIVE_WASH, stroke=C_POSITIVE, bold=True, fontsize=9)
    _box(ax, (7.0, 0.7), 3.0, 0.7,
         "Payment released\nstatus → paid · ledger updated",
         fill=C_INK, stroke=C_INK, text_color="white", bold=True, fontsize=9)
    _arrow(ax, (11.9, 3.3), (6.4, 1.05))
    _arrow(ax, (6.4, 1.05), (7.0, 1.05))

    return _save(fig, "02_procurement_flow")


# ──────────────────────────────────────────────────────────────────────────────
# 3. Sales / Revenue
# ──────────────────────────────────────────────────────────────────────────────
def fig_sales_flow():
    W, H = 14, 7
    fig, ax = _frame("3. Sales — quote to cash (inbound PO from client)", w=W, h=H)

    rows = [
        # y, items
        (5.3, [
            ("Client enquiry\nscope · site visit", C_PAPER, C_RULE, False),
            ("Quote drafted\nBoQ + GST", C_BLUE_WASH, C_BLUE, True),
            ("Quote sent", C_BLUE_WASH, C_BLUE, False),
            ("Client issues PO\n(e.g. APL/PO/26/0847)", C_ACCENT_WASH, C_ACCENT, True),
        ]),
        (3.7, [
            ("Project created\nSAB-26-NNN", C_ACCENT_WASH, C_ACCENT, True),
            ("Crew assigned\nsupervisor + team", C_POSITIVE_WASH, C_POSITIVE, False),
            ("Materials procured\n(see Procurement)", C_ACCENT_WASH, C_ACCENT, False),
            ("Work executed\ntimesheets + progress %", C_POSITIVE_WASH, C_POSITIVE, False),
        ]),
        (2.1, [
            ("Milestone reached\ne.g. 30/60/100 %", C_PAPER, C_RULE, False),
            ("Tax invoice raised\nGST split per state", C_BLUE_WASH, C_BLUE, True),
            ("Invoice sent\n(PDF + share link)", C_BLUE_WASH, C_BLUE, False),
            ("Payment received\n→ project ledger", C_POSITIVE_WASH, C_POSITIVE, True),
        ]),
    ]
    for y, items in rows:
        for i, (text, fill, stroke, bold) in enumerate(items):
            x = 0.3 + i * 3.45
            _box(ax, (x, y), 3.1, 0.85, text, fill=fill, stroke=stroke,
                 bold=bold, fontsize=9)
            if i < len(items) - 1:
                _arrow(ax, (x + 3.1, y + 0.43), (x + 3.45, y + 0.43))

    # Vertical connectors between rows
    _arrow(ax, (12.9, 5.3), (1.85, 4.55))
    _arrow(ax, (12.9, 3.7), (1.85, 2.95))

    # Footer formula
    _box(ax, (0.3, 0.4), 13.4, 0.85,
         "Project P&L  =  ∑ tax invoices   −   (∑ vendor bills + ∑ timesheet wage cost + ∑ allocated overhead)",
         fill=C_INK, stroke=C_INK, text_color="white", bold=True, fontsize=10)

    return _save(fig, "03_sales_flow")


# ──────────────────────────────────────────────────────────────────────────────
# 4. Project lifecycle + detail tabs
# ──────────────────────────────────────────────────────────────────────────────
def fig_project_lifecycle():
    W, H = 14, 7
    fig, ax = _frame("4. Project lifecycle — from kickoff to close-out", w=W, h=H)

    stages = [
        ("Planning\nBoQ + schedule", C_BLUE_WASH, C_BLUE),
        ("Kickoff\ncrew + materials", C_ACCENT_WASH, C_ACCENT),
        ("Execution\ntimesheets + GRNs", C_POSITIVE_WASH, C_POSITIVE),
        ("Monitoring\nprogress % + P&L", C_AMBER_WASH, C_AMBER),
        ("Close-out\nhandover + warranty", C_INK, C_INK),
    ]
    step_w = 2.45
    gap = 0.25
    start_x = 0.4
    for i, (text, fill, stroke) in enumerate(stages):
        x = start_x + i * (step_w + gap)
        text_color = "white" if fill == C_INK else C_INK
        _box(ax, (x, 5.3), step_w, 1.0, text, fill=fill, stroke=stroke,
             text_color=text_color, bold=True)
        if i < len(stages) - 1:
            _arrow(ax, (x + step_w, 5.8), (x + step_w + gap, 5.8))

    ax.text(0.4, 4.5, "Inside Project detail — eight tabs:",
            fontsize=10.5, fontweight="bold", color=C_INK)

    tabs = [
        ("Overview", "KPIs · timeline · docs"),
        ("P&L", "revenue − cost"),
        ("Progress", "12 milestones · % done"),
        ("Ledger", "48 entries · in/out"),
        ("Budget", "planned vs actual"),
        ("Materials", "SKU rollup · GRNs"),
        ("Invoices", "4 raised · paid status"),
        ("PO", "inbound client PO"),
    ]
    tab_w, tab_h = 3.25, 0.85
    for i, (name, sub) in enumerate(tabs):
        col = i % 4
        row = i // 4
        x = 0.4 + col * (tab_w + 0.1)
        y = 3.3 - row * 1.05
        _box(ax, (x, y), tab_w, tab_h, f"{name}\n{sub}",
             fill=C_PAPER, stroke=C_RULE, fontsize=9)

    _box(ax, (0.4, 0.4), 13.2, 0.7,
         "Every tab writes to one central project ledger — single source of truth for P&L, progress and materials rollup.",
         fill=C_ACCENT_WASH, stroke=C_ACCENT, fontsize=9)

    return _save(fig, "04_project_lifecycle")


# ──────────────────────────────────────────────────────────────────────────────
# 5. Timesheet / operations flow
# ──────────────────────────────────────────────────────────────────────────────
def fig_timesheet_flow():
    W, H = 14, 7
    fig, ax = _frame("5. Timesheets — site punch to project wage cost", w=W, h=H)

    rows = [
        (5.3, [
            ("Site crew opens\nmobile punch app", C_POSITIVE_WASH, C_POSITIVE, True),
            ("Punch in at site\nGPS-stamped", C_POSITIVE_WASH, C_POSITIVE, False),
            ("Supervisor logs hours\nper project · per skill", C_BLUE_WASH, C_BLUE, False),
            ("Timesheet entry saved\n(emp × date × project)", C_ACCENT_WASH, C_ACCENT, True),
        ]),
        (3.5, [
            ("Wage rates module\nemp × skill × rate", C_AMBER_WASH, C_AMBER, True),
            ("Hours × rate = wage cost", C_BLUE_WASH, C_BLUE, False),
            ("Posted to project ledger\nas Labour cost entry", C_ACCENT_WASH, C_ACCENT, True),
            ("Reflected in P&L\n+ progress payments", C_POSITIVE_WASH, C_POSITIVE, False),
        ]),
    ]
    for y, items in rows:
        for i, (text, fill, stroke, bold) in enumerate(items):
            x = 0.3 + i * 3.45
            _box(ax, (x, y), 3.1, 0.95, text, fill=fill, stroke=stroke,
                 bold=bold, fontsize=9)
            if i < len(items) - 1:
                _arrow(ax, (x + 3.1, y + 0.48), (x + 3.45, y + 0.48))

    _arrow(ax, (12.9, 5.3), (1.85, 4.45))

    _box(ax, (0.3, 1.4), 13.4, 1.4,
         "Overhead (rent, salaries, utilities) is allocated across projects by revenue share —\n"
         "added to each ledger every month-end so each project shows a true bottom-line P&L,\n"
         "not just direct cost.",
         fill=C_PAPER, stroke=C_RULE, fontsize=10)

    return _save(fig, "05_timesheet_flow")


# ──────────────────────────────────────────────────────────────────────────────
# 6. User management flow
# ──────────────────────────────────────────────────────────────────────────────
def fig_user_flow():
    W, H = 14, 7
    fig, ax = _frame("6. Users & roles — invite, assign, audit", w=W, h=H)

    rows = [
        (5.3, [
            ("Admin opens\nUsers & roles", C_AMBER_WASH, C_AMBER, True),
            ("Invite user\nname + email + role", C_ACCENT_WASH, C_ACCENT, True),
            ("User entry created\nstatus = invited", C_BLUE_WASH, C_BLUE, False),
            ("Email sent\n(SSO set-up link)", C_BLUE_WASH, C_BLUE, False),
        ]),
        (3.7, [
            ("Role = Admin / Mgr /\nSupervisor / Crew / View", C_PAPER, C_RULE, False),
            ("Role matrix grants\nmodule permissions", C_AMBER_WASH, C_AMBER, False),
            ("Assign to projects\n(one-to-many)", C_ACCENT_WASH, C_ACCENT, True),
            ("User signs in →\nsees only assigned", C_POSITIVE_WASH, C_POSITIVE, True),
        ]),
        (2.0, [
            ("Activity log\nper user · per action", C_PAPER, C_RULE, False),
            ("Force password reset", C_AMBER_WASH, C_AMBER, False),
            ("Suspend / re-enable", C_AMBER_WASH, C_AMBER, False),
            ("Remove user\n(audit preserved)", C_ALERT_WASH, C_ALERT, False),
        ]),
    ]
    for y, items in rows:
        for i, (text, fill, stroke, bold) in enumerate(items):
            x = 0.3 + i * 3.45
            _box(ax, (x, y), 3.1, 0.95, text, fill=fill, stroke=stroke,
                 bold=bold, fontsize=9)
            if i < len(items) - 1:
                _arrow(ax, (x + 3.1, y + 0.48), (x + 3.45, y + 0.48))

    _arrow(ax, (12.9, 5.3), (1.85, 4.65))
    _arrow(ax, (12.9, 3.7), (1.85, 2.95))

    _box(ax, (0.3, 0.4), 13.4, 0.95,
         "Every admin action — role edit, project assign, suspend, remove — is appended to the user's activity log.\n"
         "Logs are never deleted (compliance trail).",
         fill=C_INK, stroke=C_INK, text_color="white", fontsize=9.5)

    return _save(fig, "06_user_flow")


# ──────────────────────────────────────────────────────────────────────────────
# 7. PO approval threshold decision
# ──────────────────────────────────────────────────────────────────────────────
def fig_po_approval():
    W, H = 11, 8
    fig, ax = _frame("7. PO approval — threshold decision", w=W, h=H)

    _box(ax, (4.0, 6.6), 3.0, 0.8, "Draft PO\ntotal = subtotal + GST",
         fill=C_PAPER, stroke=C_RULE, bold=True)

    _diamond(ax, (3.7, 4.7), 3.6, 1.5, "total ≤ ₹1L ?",
             fill=C_AMBER_WASH, stroke=C_AMBER, fontsize=10)
    _arrow(ax, (5.5, 6.6), (5.5, 6.2))

    _box(ax, (0.3, 5.05), 3.0, 0.8, "YES → auto-approved\nstatus jumps to 'sent'",
         fill=C_POSITIVE_WASH, stroke=C_POSITIVE, bold=True, fontsize=9)
    _arrow(ax, (3.7, 5.45), (3.3, 5.45), label="yes")

    _diamond(ax, (3.7, 2.6), 3.6, 1.5, "total ≤ ₹10L ?",
             fill=C_AMBER_WASH, stroke=C_AMBER, fontsize=10)
    _arrow(ax, (5.5, 4.7), (5.5, 4.1), label="no")

    _box(ax, (0.3, 2.95), 3.0, 0.8, "YES → PM approval\n(Anita Rao)",
         fill=C_BLUE_WASH, stroke=C_BLUE, bold=True, fontsize=9)
    _arrow(ax, (3.7, 3.35), (3.3, 3.35), label="yes")

    _box(ax, (4.0, 0.8), 3.0, 0.8, "NO → Director approval",
         fill=C_ALERT_WASH, stroke=C_ALERT, bold=True)
    _arrow(ax, (5.5, 2.6), (5.5, 1.6), label="no")

    _box(ax, (7.7, 3.7), 3.0, 1.4,
         "Approved POs land\nat status = sent\n\nPending POs wait\non the approver",
         fill=C_PAPER, stroke=C_RULE, fontsize=9.5)

    return _save(fig, "07_po_approval")


# ──────────────────────────────────────────────────────────────────────────────
# 8. Vendor bill state machine
# ──────────────────────────────────────────────────────────────────────────────
def fig_bill_state_machine():
    W, H = 14, 7
    fig, ax = _frame("8. Vendor bill state machine (3-way match)", w=W, h=H)

    states = {
        "draft":        (0.5, 5.0, C_PAPER, C_RULE),
        "pending-match":(3.6, 5.0, C_AMBER_WASH, C_AMBER),
        "matched":      (6.7, 5.0, C_BLUE_WASH, C_BLUE),
        "approved":     (9.8, 5.0, C_ACCENT_WASH, C_ACCENT),
        "discrepancy":  (3.6, 3.2, C_ALERT_WASH, C_ALERT),
        "overdue":      (3.6, 1.4, C_ALERT_WASH, C_ALERT),
        "paid":         (9.8, 1.4, C_POSITIVE_WASH, C_POSITIVE),
    }
    box_w, box_h = 3.0, 0.8
    for name, (x, y, fill, stroke) in states.items():
        _box(ax, (x, y), box_w, box_h, name, fill=fill, stroke=stroke,
             bold=True, fontsize=10)

    def edge(a, b, label=None):
        ax_, ay_, *_ = states[a]
        bx_, by_, *_ = states[b]
        # default: from right edge of a to left edge of b
        sx, sy = ax_ + box_w, ay_ + box_h / 2
        ex, ey = bx_, by_ + box_h / 2
        # if same x: go vertical
        if abs(ax_ - bx_) < 0.01:
            sx = ax_ + box_w / 2
            ex = bx_ + box_w / 2
            sy = ay_ if ay_ > by_ else ay_ + box_h
            ey = by_ + box_h if ay_ > by_ else by_
        _arrow(ax, (sx, sy), (ex, ey), label=label)

    edge("draft", "pending-match", "record")
    edge("pending-match", "matched", "match ok")
    edge("matched", "approved", "schedule")
    edge("pending-match", "discrepancy", "mismatch")
    edge("discrepancy", "matched", "resolve")
    edge("pending-match", "overdue", "due passed")
    edge("approved", "paid", "release")
    # paid wraps under from approved (already covered above)

    _box(ax, (0.3, 0.2), 13.4, 0.8,
         "3-way match rule: PO total  ≈  ∑(accepted GRN lines × PO rate)  ≈  vendor bill total   (tolerance < ₹100)",
         fill=C_INK, stroke=C_INK, text_color="white", fontsize=9.5)

    return _save(fig, "08_bill_state_machine")


# ──────────────────────────────────────────────────────────────────────────────
# Build the Word document
# ──────────────────────────────────────────────────────────────────────────────
def build_document(image_paths: dict):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Cover
    title = doc.add_paragraph()
    run = title.add_run("SAB India Tracker")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2B, 0x26, 0x24)

    sub = doc.add_paragraph()
    run = sub.add_run("End-to-end activity flow charts")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xD4, 0x71, 0x2B)

    meta = doc.add_paragraph()
    run = meta.add_run(
        "Fire, plumbing and utility installation operations · internal tool · v3.2\n"
        "Covers sign-in → sales → project execution → procurement → invoicing → reporting.\n"
        "Each diagram is a standalone view; flows connect through a single project ledger."
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x4F, 0x4A)

    doc.add_paragraph()

    toc_head = doc.add_paragraph()
    run = toc_head.add_run("Contents")
    run.bold = True
    run.font.size = Pt(14)

    toc_entries = [
        "1. System navigation map",
        "2. Procurement — vendor to payment",
        "3. Sales — quote to cash",
        "4. Project lifecycle",
        "5. Timesheets & operations",
        "6. Users & roles",
        "7. PO approval thresholds",
        "8. Vendor bill state machine",
    ]
    for entry in toc_entries:
        p = doc.add_paragraph(entry, style="List Number")
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    sections = [
        (
            "1. System navigation map",
            image_paths["01_nav_map"],
            "The app has four left-rail sections — Sales, Operations, Procurement, Admin — "
            "plus two side-channels: the mobile punch app for site crew and the public invoice "
            "share page for clients. Sign-in lands on the Operations Home, which surfaces "
            "action items pulled from every other module.",
        ),
        (
            "2. Procurement — vendor to payment",
            image_paths["02_procurement_flow"],
            "The most ledger-sensitive path in the app. Every PO draft passes a three-way "
            "threshold test before it can be sent to the vendor. The GRN captures shortfall and "
            "rejection data that drives vendor credit notes. Bills then reconcile through a "
            "3-way match (PO ↔ GRN ↔ Bill) before they are routed to finance for payment "
            "scheduling.",
        ),
        (
            "3. Sales — quote to cash",
            image_paths["03_sales_flow"],
            "Sales starts with a client enquiry and ends with a payment posted to the project "
            "ledger. Once the client issues their inbound PO, a project is created and the "
            "execution track picks up — crew assignment, materials procurement (separate flow, "
            "diagram 2), and milestone-based invoicing. Every invoice writes to the same "
            "project ledger where vendor bills, wages and overhead also land, producing one "
            "P&L per project.",
        ),
        (
            "4. Project lifecycle",
            image_paths["04_project_lifecycle"],
            "Inside the Project Detail view, work is organised across eight tabs. Each tab "
            "writes back to the same central project ledger — that is why P&L, progress and "
            "materials rollup always stay in sync. Close-out captures handover docs and starts "
            "the warranty clock.",
        ),
        (
            "5. Timesheets & operations",
            image_paths["05_timesheet_flow"],
            "Site crew punches in via the mobile app (GPS-stamped). Supervisors log hours per "
            "project per skill; the Wage rates module provides the multiplier that turns hours "
            "into labour cost. That cost lands on the project ledger as a Labour entry, feeding "
            "P&L, progress payments and overhead allocation.",
        ),
        (
            "6. Users & roles",
            image_paths["06_user_flow"],
            "Admins invite users by email and role. The role matrix grants module-level "
            "permissions (Admin / Manager / Supervisor / Crew / View-only). Project assignment "
            "is separate and narrows what a given user actually sees once signed in. Every "
            "action — role edit, project assign, suspend, remove — is appended to the user's "
            "activity log for compliance.",
        ),
        (
            "7. PO approval thresholds",
            image_paths["07_po_approval"],
            "A single helper (approvalTierFor) drives the whole approval routing. Under ₹1L "
            "the PO is auto-approved and immediately marked 'sent'. Between ₹1L and ₹10L the "
            "PO waits for a project manager's click. Over ₹10L only a Director can sign off. "
            "The tier is stored with the PO so an auditor can later see why a given PO went "
            "where it did.",
        ),
        (
            "8. Vendor bill state machine",
            image_paths["08_bill_state_machine"],
            "Every vendor bill moves through a fixed set of states. The green path — draft → "
            "pending-match → matched → approved → paid — is the happy flow. Bills can sidetrack "
            "through 'discrepancy' (resolved by the buyer) or 'overdue' (passed due date "
            "without payment). Tolerance for the 3-way match is currently ₹100.",
        ),
    ]

    for heading, img_path, description in sections:
        h = doc.add_paragraph()
        run = h.add_run(heading)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2B, 0x26, 0x24)

        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic.add_run()
        run.add_picture(str(img_path), width=Inches(6.8))

        body = doc.add_paragraph(description)
        body.paragraph_format.space_after = Pt(14)

        doc.add_page_break()

    h = doc.add_paragraph()
    run = h.add_run("Legend")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2B, 0x26, 0x24)

    legend = [
        ("Orange box", "Procurement or core record (vendor, PO, bill, project)"),
        ("Blue box", "Sales-side record (quote, invoice, payment)"),
        ("Green box", "Successful completion / happy path"),
        ("Amber diamond", "Decision point (threshold, match, status check)"),
        ("Red box", "Failure, discrepancy or overdue state"),
        ("Dark box", "System-level or cross-cutting fact (e.g. P&L formula)"),
    ]
    for label, meaning in legend:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}  ")
        run.bold = True
        p.add_run(f"— {meaning}")

    out_path = DOCS_DIR / "SAB_India_Tracker_Flowcharts.docx"
    doc.save(out_path)
    return out_path


def main():
    image_paths = {}
    image_paths["01_nav_map"] = fig_nav_map()
    image_paths["02_procurement_flow"] = fig_procurement_flow()
    image_paths["03_sales_flow"] = fig_sales_flow()
    image_paths["04_project_lifecycle"] = fig_project_lifecycle()
    image_paths["05_timesheet_flow"] = fig_timesheet_flow()
    image_paths["06_user_flow"] = fig_user_flow()
    image_paths["07_po_approval"] = fig_po_approval()
    image_paths["08_bill_state_machine"] = fig_bill_state_machine()

    out = build_document(image_paths)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
